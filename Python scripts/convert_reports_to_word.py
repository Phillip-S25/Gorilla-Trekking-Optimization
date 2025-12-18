"""
Convert all TXT reports to Word documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Define reports folder
reports_folder = "../Reports"

# List of reports to convert
reports = [
    'eda_summary.txt',
    'elasticity_results.txt',
    'optimization_results.txt',
    'EXECUTIVE_SUMMARY.txt'
]

print("="*70)
print("CONVERTING TXT REPORTS TO WORD DOCUMENTS")
print("="*70)

for report_file in reports:
    input_path = os.path.join(reports_folder, report_file)
    output_path = os.path.join(reports_folder, report_file.replace('.txt', '.docx'))
    
    print(f"\nProcessing: {report_file}")
    
    try:
        # Read the text file
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create a new Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Parse content and add to document
        lines = content.split('\n')
        
        for line in lines:
            # Check if line is a header (contains === or ---)
            if '=' * 20 in line:
                # Major heading
                if lines.index(line) > 0:
                    para = doc.add_paragraph()  # Add space before
                continue
            elif '-' * 20 in line:
                # Section divider
                continue
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
                continue
            else:
                # Check if it's a title line (next line has ===)
                idx = lines.index(line)
                is_title = (idx + 1 < len(lines) and '=' * 20 in lines[idx + 1])
                is_heading = (idx + 1 < len(lines) and '-' * 20 in lines[idx + 1])
                
                if is_title:
                    # Main title
                    para = doc.add_heading(line.strip(), level=0)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif is_heading:
                    # Section heading
                    para = doc.add_heading(line.strip(), level=1)
                elif line.startswith('  •') or line.startswith('•'):
                    # Bullet point
                    para = doc.add_paragraph(line.strip().lstrip('•').strip(), style='List Bullet')
                elif line.startswith('  -') or line.startswith('-'):
                    # Bullet point
                    para = doc.add_paragraph(line.strip().lstrip('-').strip(), style='List Bullet')
                elif ':' in line and len(line) < 80:
                    # Key-value pair (possibly bold the key)
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        para = doc.add_paragraph()
                        para.add_run(parts[0] + ':').bold = True
                        para.add_run(parts[1])
                    else:
                        para = doc.add_paragraph(line)
                else:
                    # Regular paragraph
                    para = doc.add_paragraph(line)
        
        # Save the document
        doc.save(output_path)
        print(f"✓ Saved: {report_file.replace('.txt', '.docx')}")
        
    except Exception as e:
        print(f"✗ Error converting {report_file}: {e}")

print("\n" + "="*70)
print("CONVERSION COMPLETE")
print("="*70)
print(f"\nWord documents saved in: {reports_folder}/")
